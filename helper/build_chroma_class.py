import os, re, csv, json, uuid, glob
import html, math
import pickle

from pathlib import Path
from collections import defaultdict, Counter
from typing import List, Tuple, Dict, Optional, Any
from chromadb import PersistentClient
from ollama import Client
from pypdf import PdfReader

import chromadb
from chromadb.config import Settings
import pypdf
import pandas as pd
from docx import Document

from rank_bm25 import BM25Okapi

try:
    import tabula
    USE_TABULA = True
except ImportError:
    USE_TABULA = False


class RAGSystem:
    """En klass för att hantera RAG-flöden med Ollama och ChromaDB.
    I/O Vector databas hantering
    """
    def __init__(self,
                 embed_model        = None  ,
                 llm_model          = None  ,
                 chroma_dir         = None  ,
                 data_root          = None  ,
                 whoosh_dir         = None  ,
                 collection_name    = None  ,
                 ollama_host        = None  ,
                 batch_size         = None  ,
                 max_workers        = None  ,
                 chunk_size         = None  ,
                 overlap            = None  ,
                 questions          = None  ,
                 summary_prompt     = None  ,
                 bCreate            = False ):
        """
        Initiera RAGSystem.

        Args:
            embedding_model: Modell för embeddings (Ollama).
            llm_model: Språkmodell för generering.
            chroma_path: Sökväg till ChromaDB‑katalogen. Om None används in‑memory index.
            collection_name: Namn på ChromaDB‑samlingen.
        """
        if data_root is None :
            self.DATA_ROOT          = Path(os.getenv("DATA_ROOT"    , "./data"))
            self.data_root          = self.DATA_ROOT
        else :
            self.DATA_ROOT          = Path(data_root)
            self.data_root          = Path(data_root)

        if chroma_dir is None :
            self.CHROMA_DIR         = os.getenv("CHROMA_DIR"        , "./chroma_db")
            self.chroma_path        = self.CHROMA_DIR
        else :
            self.CHROMA_DIR         = chroma_dir
            self.chroma_path        = chroma_dir

        self.BM25_DIR = f"{self.chroma_path}/bm25.pkl"

        if whoosh_dir is None :
            self.WHOOSH_DIR         = os.getenv("WHOOSH_DIR"        , "./chroma_db/whoosh_index")
            self.whoosh_dir         = self.WHOOSH_DIR
        else:
            self.WHOOSH_DIR         = whoosh_dir
            self.whoosh_dir         = whoosh_dir

        if collection_name is None :
            self.COLLECTION_NAME    = os.getenv("COLLECTION_NAME"   , "spectrum_data")
            self.collection_name    = self.COLLECTION_NAME
        else:
            self.COLLECTION_NAME    = collection_name
            self.collection_name    = collection_name

        if ollama_host is None :
            self.OLLAMA_HOST        = os.getenv("OLLAMA_HOST"       , "http://127.0.0.1:11434")
            self.ollama_host        = self.OLLAMA_HOST
        else:
            self.OLLAMA_HOST        = ollama_host
            self.ollama_host        = ollama_host

        if embed_model is None :
            self.EMBED_MODEL        = os.getenv("EMBED_MODEL"       , "nomic-embed-text-v2-moe")
            self.embed_model        = self.EMBED_MODEL
        else:
            self.EMBED_MODEL        = embed_model
            self.embed_model        = embed_model

        if llm_model is None :
            self.LLM_MODEL          = os.getenv("LLM_MODEL"         , "qwen3:30b-a3b")
            self.llm_model          = self.LLM_MODEL
        else:
            self.LLM_MODEL          = llm_model
            self.llm_model          = llm_model

        if batch_size is None :
            self.BATCH_SIZE         = 32
            self.batch_size         = self.BATCH_SIZE
        else:
            self.BATCH_SIZE         = batch_size
            self.batch_size         = batch_size

        if max_workers is None :
            self.MAX_WORKERS        = 4
            self.max_workers        = self.MAX_WORKERS
        else:
            self.MAX_WORKERS        = max_workers
            self.max_workers        = max_workers

        if chunk_size is None :
            self.chunk_size_        = 2000
        else :
            self.chunk_size_        = chunk_size

        if overlap is None:
            self.overlap_           = 200
        else :
            self.overlap_           = overlap
        self.MIN_CHARS              = 5
        self.max_length_            = self.chunk_size_
        self.default_embedding_length_ = 768

        self.ids        = None
        self.docs       = None
        self.metas      = None
        self.emb_texts  = None
        self.all_freqs  = None
        self.all_terms  = None
        self.texts      = None

        self.STOPWORDS = {
            "och","eller","med","för","som","den","det","att","av","till","i","på","är","en","ett","sig","om","under","efter","över","vid",
            "khz","mhz","ghz","hz","db","kw","w","v","a","nan","none","null"
        }

        self.embedding_model = self.embed_model

        # Behåll stöd för in‑memory index (fallback)
        self.chunks: List[str] = []
        self.embeddings: List[List[float]] = []

        if not questions is None :
            self.questions_ = questions

        if not summary_prompt is None :
            self.summary_prompt_ = summary_prompt

        self.USE_TABULA = False

        self.id_to_index = {}   # Chroma ID -> BM25 index
        self.index_to_id = {}   # BM25 index -> Chroma ID
        self.id_to_doc = {}     # ID -> full doc (text + meta)

        # Initiera ChromaDB om en sökväg angivits
        self.chroma_client  = None
        self.collection     = None
        if self.chroma_path:
            self.init_chroma()
            self.load_bm25()

    def is_good_text(self, t: str) -> bool:
        if not t or len(t) < 50:
            return False
        words = t.split()
        if len(set(words)) < 5:
            return False
        return True

    def split_semantic(self, text: str) -> List[str]:
        import re
        blocks = re.split(r"\n\s*\n", text)
        return [b.strip() for b in blocks if len(b.strip()) > 50]

    def adaptive_chunk(self, texts: List[str], max_tokens=400) -> List[str]:
        chunks = []
        current = []
        current_len = 0

        for t in texts:
            l = len(t.split())
            if current_len + l > max_tokens and current:
                chunks.append(" ".join(current))
                current = []
                current_len = 0

            current.append(t)
            current_len += l

        if current:
            chunks.append(" ".join(current))

        return chunks

    def build_embedding_text(self, text: str, meta: dict) -> str:
        return f"""
    TYPE: {meta.get('type')}

    FREQUENCY_FROM_MHZ: {meta.get('frequency_from_mhz')}
    FREQUENCY_TO_MHZ: {meta.get('frequency_to_mhz')}
    FREQ_CENTER: {meta.get('freq_center')}
    FREQ_WIDTH: {meta.get('freq_width')}

    CONTENT:
    {text}
    """.strip()

    def deduplicate(self, texts, metas, ids):
        seen = set()
        new_t, new_m, new_i = [], [], []

        for t, m, i in zip(texts, metas, ids):
            h = hash(t)
            if h in seen:
                continue
            seen.add(h)
            new_t.append(t)
            new_m.append(m)
            new_i.append(i)

        return new_t, new_m, new_i

    def generate_embeddings_batch(self, texts):
        try:
            response = self.ollama.embed(
                model=self.embedding_model,
                input=texts
            )
            return response.embeddings
        except Exception as e:
            print(f"Batch embedding error: {e}")
            return [[0.0]*self.default_embedding_length_] * len(texts)

    def init_chroma(self):
        """Skapa eller anslut till ChromaDB‑samling."""
        os.makedirs(self.chroma_path, exist_ok=True)
        self.ollama         = Client(host=self.OLLAMA_HOST)
        self.chroma_client  = chromadb.PersistentClient(path=self.chroma_path)
        # Försök hämta befintlig samling, annars skapa ny
        try:
            self.collection = self.chroma_client.get_collection(self.collection_name)
            print(f"Ansluten till befintlig ChromaDB‑samling '{self.collection_name}'")
        except Exception:
            self.collection = self.chroma_client.create_collection(self.collection_name)
            print(f"Skapade ny ChromaDB‑samling '{self.collection_name}'")
        self.chroma = self.chroma_client

        # ------------------------ FILHANTERING (befintliga metoder) ------------------------
        self.chunks: List[str] = []
        self.embeddings: List[List[float]] = []
        self.default_embedding_length_ = 768

    def read_pdf(self, filepath):
        """Läser innehållet från en PDF-fil och returnerar det som en sträng."""
        import pypdf
        text = ""
        try:
            with open(filepath, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text
        except Exception as e:
            print(f"Fel vid läsning av PDF {filepath}: {e}")
            return ""

    def generate_embedding(self, text):
        """Genererar embedding med Ollama."""
        try:
            response = self.ollama.embed(model=self.embedding_model, input=text)
            if response and hasattr(response, 'embeddings') and len(response.embeddings) > 0:
                return response.embeddings[0] if isinstance(response.embeddings[0], list) else response.embeddings
            else:
                return [0.0] * self.default_embedding_length_
        except Exception as e:
            print(f"Embedding misslyckades: {e}")
            return [0.0] * self.default_embedding_length_

    # ------------------------ VERKTYG FÖR BM25 ---------------------------
    def tokenize(self, text):
        import re
        text = text.lower()
        text = re.sub(r"[^\w\s\.\-]", " ", text)
        return text.split()


    def build_and_save_bm25(self, texts, metas, ids):
        from rank_bm25 import BM25Okapi
        import pickle

        print("Building BM25 index...")

        tokenized = [self.tokenize(t) for t in texts]
        bm25 = BM25Okapi(tokenized)

        id_to_index = {}
        index_to_id = {}
        id_to_doc = {}

        for i, (doc_id, text, meta) in enumerate(zip(ids, texts, metas)):
            id_to_index[doc_id] = i
            index_to_id[i] = doc_id
            id_to_doc[doc_id] = {
                "text": text,
                "meta": meta
            }

        self.tokenized_corpus = tokenized
        self.bm25        = bm25
        self.id_to_index = id_to_index
        self.index_to_id = index_to_id
        self.id_to_doc   = id_to_doc

        with open(self.BM25_DIR, "wb") as f:
            pickle.dump({
                "bm25": bm25,
                "id_to_index": id_to_index,
                "index_to_id": index_to_id,
                "id_to_doc": id_to_doc
            }, f)

        print(f"✅ BM25 saved ({len(texts)} docs)")


    def load_bm25(self):
        import pickle

        path = self.BM25_DIR

        if not os.path.exists(path):
            print("⚠️ No BM25 index found")
            return

        with open(path, "rb") as f:
            data = pickle.load(f)

        self.bm25        = data["bm25"]
        self.id_to_index = data["id_to_index"]
        self.index_to_id = data["index_to_id"]
        self.id_to_doc   = data["id_to_doc"]

        print(f"✅ BM25 loaded ({len(self.id_to_doc)} docs)")


    def bm25_query(self, query, k=30):
        tokens = self.tokenize(query)
        scores = self.bm25.get_scores(tokens)

        top_idx = sorted(
            range(len(scores)),
            key=lambda i: scores[i],
            reverse=True
        )[:k]

        results = []
        for i in top_idx:
            doc_id = self.index_to_id[i]
            doc = self.id_to_doc[doc_id]

            results.append({
                "id": doc_id,
                "text": doc["text"],
                "meta": doc["meta"],
                "score": scores[i],
                "source": "bm25"
            })

        return results


    # ------------------------ CHROMADB-INDEXERING ------------------------
    def add_documents(self, texts, metadatas=None, ids=None):

        if not self.collection:
            raise RuntimeError("ChromaDB ej initierad")

        # Deduplicera
        texts, metadatas, ids = self.deduplicate(texts, metadatas, ids)

        # Filtrera
        filtered = [(t, m, i) for t, m, i in zip(texts, metadatas, ids) if self.is_good_text(t)]
        if not filtered:
            return

        texts, metadatas, ids = zip(*filtered)
        texts = list(texts)
        metadatas = list(metadatas)
        ids = list(ids)

        batch_size = 100

        for i in range(0, len(texts), batch_size):

            batch_texts = texts[i:i+batch_size]
            batch_meta = metadatas[i:i+batch_size]
            batch_ids = ids[i:i+batch_size]

            # Bygg embedding-text PER BATCH
            texts_structured = [
                self.build_embedding_text(t, m)
                for t, m in zip(batch_texts, batch_meta)
            ]

            # Generera embeddings PER BATCH
            embeddings = self.generate_embeddings_batch(texts_structured)

            # KRITISK CHECK
            if len(embeddings) != len(batch_texts):
                print("❌ Embedding mismatch!")
                print(f"texts: {len(batch_texts)} embeddings: {len(embeddings)}")
                continue

            self.collection.add(
                documents=batch_texts,
                embeddings=embeddings,
                metadatas=batch_meta,
                ids=batch_ids
            )

    def add_documents_leg(self, texts, metadatas=None, ids=None):
        if not self.collection:
            raise RuntimeError("ChromaDB ej initierad")

        # Deduplicera
        texts, metadatas, ids = self.deduplicate(texts, metadatas, ids)

        # Filtrera lågkvalitativ text
        filtered = [(t, m, i) for t, m, i in zip(texts, metadatas, ids) if self.is_good_text(t)]
        if not filtered:
            return

        texts, metadatas, ids = zip(*filtered)

        # Bygg strukturerad embedding-text
        texts_structured = [
            self.build_embedding_text(t, m) for t, m in zip(texts, metadatas)
        ]

        # Batch embeddings
        embeddings = self.generate_embeddings_batch(list(texts_structured))

        # Batch insert
        batch_size = 100
        for i in range(0, len(texts), batch_size):
            self.collection.add(
                documents=texts[i:i+batch_size],
                embeddings=embeddings[i:i+batch_size],
                metadatas=metadatas[i:i+batch_size],
                ids=ids[i:i+batch_size]
            )

    def add_files(self, directory, extensions = ['.pdf', '.xlsx', '.docx', '.csv', '.json']):
        """
        Läser alla filer med givna ändelser i en katalog (rekursivt) och lägger till dem i ChromaDB.
        """
        all_texts = []
        all_metadatas = []
        all_ids = []

        for ext in extensions:
            for filepath in glob.glob(os.path.join(directory, '**', f'*{ext}'), recursive=True):
                print(f"Bearbetar {filepath} ...")
                text = self.read_any_file(filepath)
                if not text:
                    continue
                # Dela upp i chunks
                chunks = self.chunk_text(text)
                # Metadata för varje chunk
                base_meta = {"source": filepath, "type": ext.lstrip('.')}
                for i, chunk in enumerate(chunks):
                    all_texts.append(chunk)
                    meta = base_meta.copy()
                    meta["chunk_index"] = i
                    all_metadatas.append(meta)
                    # Skapa ett ID baserat på fil och chunk
                    import hashlib
                    id_str = f"{filepath}_{i}"
                    all_ids.append(hashlib.md5(id_str.encode()).hexdigest())

        # Lägg till alla i databasen
        if all_texts:
            self.add_documents(all_texts, all_metadatas, all_ids)
            self.build_and_save_bm25(all_texts, all_metadatas, all_ids)
            print(f"Indexerat {len(all_texts)} chunks från {directory}")
        else:
            print("Inga dokument hittades.")

    def create_chromadb_from_data(self, table_as_document: bool = True,
                                  chunk_tables: bool = False , group_tables: bool = True,
                                  group_token_limit: int = 400 ):
        """
        Skapar ChromaDB‑databasen genom att läsa alla filer i DATA_ROOT.
        Tabellinformation (rader) sparas som separata dokument för att inte gå förlorad.

        Args:
            table_as_document: Om True sparas varje tabellrad som ett eget dokument.
                            Annars stoppas hela tabellens text in i ett vanligt chunk.
            chunk_tables:   Om True delas även tabellrader i mindre bitar (rekommenderas False).
        """
        if not self.collection:
            self.init_chroma()

        all_texts = []
        all_metadatas = []
        all_ids = []

        # Gå igenom alla filer
        for file_path in self.DATA_ROOT.rglob("*"):
            if not file_path.is_file():
                continue
            ext = file_path.suffix.lower()
            if ext not in ['.pdf', '.csv', '.xlsx', '.docx', '.json']:
                continue

            print(f"📄 Bearbetar: {file_path.name}")
            items = []
            # Hantera olika filtyper med fokus på tabeller
            if ext == '.csv' and table_as_document:
                items = self._extract_csv_rows(file_path)
            elif ext == '.xlsx' and table_as_document:
                items = self._extract_excel_rows(file_path)
            elif ext == '.pdf' and table_as_document:
                items = self._extract_pdf_tables_and_text(file_path)
            elif ext == '.docx' and table_as_document:
                items = self._extract_docx_tables_and_text(file_path)
            elif ext == '.json':
                items = self._extract_json_items(file_path)
            else:
                print ( "VARNING: Fallback, läs som vanlig text" )
                text = self.read_any_file(file_path)
                if isinstance(text, str) and text.strip():
                    #items = [(text, {"type": "text", "source": str(file_path)})]
                    blocks = self.split_semantic(text)
                    chunks = self.adaptive_chunk(blocks, max_tokens=400)

                    items = []
                    for i, chunk in enumerate(chunks):
                        meta = {
                            "type": "text_chunk",
                            "source": str(file_path),
                            "chunk_index": i
                        }
                        items.append((chunk, meta))
                else:
                    continue

            if group_tables :
                texts = [t for t, _ in items]
                metas = [m for _, m in items]

                grouped_texts = self.adaptive_chunk(texts, max_tokens=group_token_limit)

                grouped_items = []
                for i, gtext in enumerate(grouped_texts):
                    meta = {
                        "type": "grouped_block",
                        "source": str(file_path),
                        "block_index": i
                    }
                    grouped_items.append((gtext, meta))
                items = grouped_items
                print(f"  📊 Skapade {len(items)} representationer från {len(texts)}")


            # Lägg till varje utvunnet objekt
            for text, meta in items:
                if not text or len(text) < 5:
                    continue
                if chunk_tables:
                    chunks = self.chunk_text(text)
                    for i, chunk in enumerate(chunks):
                        all_texts.append(chunk)
                        m = meta.copy()
                        m["chunk_index"] = i
                        all_metadatas.append(m)
                        all_ids.append(str(uuid.uuid4()))
                else:
                    all_texts.append(text)
                    all_metadatas.append(meta)
                    all_ids.append(str(uuid.uuid4()))

        # Lägg till alla dokument i ChromaDB (med batcher)
        if all_texts:
            self.add_documents(all_texts, all_metadatas, all_ids)
            self.build_and_save_bm25(all_texts, all_metadatas, all_ids)
            print(f"\n✅ Indexerat {len(all_texts)} dokument (rader + textstycken) i '{self.collection_name}'")
        else:
            print("❌ Inga dokument hittades att indexera.")

    # ------------------- Hjälpmetoder för tablextraktion -------------------
    def _extract_csv_rows(self, file_path):
        import pandas as pd
        items = []

        # 1. Läs CSV-filen robust (autodetect separator)
        separators = [';', ',', '\t']
        df = None
        for sep in separators:
            try:
                df = pd.read_csv(
                    file_path,
                    sep=sep,
                    dtype=str,
                    keep_default_na=False,
                    on_bad_lines='skip',
                    encoding='utf-8'
                )
                if df.shape[1] > 1:
                    break
            except:
                continue
        if df is None or df.empty:
            print(f"  ⚠️ Kunde inte läsa CSV: {file_path}")
            return items

        # Normalisera kolumnnamn (gör om till gemener, ta bort mellanslag)
        df.columns = [col.strip().lower().replace(' ', '_') for col in df.columns]

        # 2. Identifiera typ baserat på kolumnnamn
        if 'allokerat_band_enligt_itu-rr' in df.columns or 'allokerade_tjänster_enligt_itu-rr' in df.columns:
            # ----- ITU-typ (med pipe-separerade användningar) -----
            return self._extract_itu_csv_rows(df, file_path)
        elif 'från_(mhz)' in df.columns and 'till_(mhz)' in df.columns:
            # ----- Inriktningsplantyp -----
            return self._extract_plan_csv_rows(df, file_path)
        else:
            # ----- Okänd typ – fallback: varje rad som ett dokument -----
            print(f"  ⚠️ Okänd CSV-struktur i {file_path.name}, använder fallback (en rad = ett dokument)")
            return self._extract_generic_csv_rows(df, file_path)

    # ------------------- Inriktningsplan-typ -------------------
    def _extract_plan_csv_rows_sugga(self, df, file_path):
        items = []
        rows_text = []
        metas = []

        for idx, row in df.iterrows():
            freq_from = row.get('från_(mhz)', '').strip()
            freq_to = row.get('till_(mhz)', '').strip()

            if not freq_from or not freq_to:
                continue

            try:
                f1 = float(freq_from.replace(',', '.'))
                f2 = float(freq_to.replace(',', '.'))
            except:
                f1, f2 = None, None

            usage = str(row.get('användning_idag', '')).strip()

            text = f"Frekvens: {freq_from}-{freq_to} MHz | Användning: {usage}"

            meta = {
                "type": "allocation_plan",
                "source": str(file_path),
                "row": idx,
                "frequency_from_mhz": f1,
                "frequency_to_mhz": f2,
            }

            if f1 and f2:
                meta["freq_center"] = (f1 + f2) / 2
                meta["freq_width"] = (f2 - f1)

            rows_text.append(text)
            metas.append(meta)

        # GROUPING
        grouped_texts = self.adaptive_chunk(rows_text, max_tokens=400)

        for i, gtext in enumerate(grouped_texts):
            meta = {
                "type": "allocation_plan_block",
                "source": str(file_path),
                "block_index": i
            }
            items.append((gtext, meta))

        print(f"  📊 Skapade {len(items)} representationer från plan {file_path.name}")
        return items

    def _extract_plan_csv_rows(self, df, file_path):
        items       = []
        rows_text   = []
        metas       = []
        # Mappa kolumner (använd exakta namn efter normalisering)
        col_from = 'från_(mhz)'
        col_to = 'till_(mhz)'
        col_mangd = 'mängd_(mhz)' if 'mängd_(mhz)' in df.columns else None
        col_popular_name = 'populärnamn' if 'populärnamn' in df.columns else None
        col_duplex = 'duplexband' if 'duplexband' in df.columns else None
        col_usage_today = 'användning_idag' if 'användning_idag' in df.columns else None
        col_allocation_form = 'tilldelningsform' if 'tilldelningsform' in df.columns else None
        col_planned_change = 'planerad_förändring' if 'planerad_förändring' in df.columns else None

        for idx, row in df.iterrows():
            freq_from = row.get(col_from, '').strip()
            freq_to = row.get(col_to, '').strip()
            if not freq_from or not freq_to:
                continue

            # Rensa bort eventuella icke-numeriska tecken (men behåll som sträng för sökning)
            # Vi behåller originalsträngen men kan också skapa numeriska värden för filtrering
            try:
                freq_from_num = float(freq_from.replace(',', '.'))
                freq_to_num = float(freq_to.replace(',', '.'))
            except:
                freq_from_num = None
                freq_to_num = None

            mangd = row.get(col_mangd, '').strip() if col_mangd else ''
            popular_name = row.get(col_popular_name, '').strip() if col_popular_name else ''
            duplex = row.get(col_duplex, '').strip() if col_duplex else ''
            usage_raw = row.get(col_usage_today, '').strip() if col_usage_today else ''
            allocation_form = row.get(col_allocation_form, '').strip() if col_allocation_form else ''
            planned_change = row.get(col_planned_change, '').strip() if col_planned_change else ''

            # Rensa användningstexten: ersätt radbrytningar med mellanslag och komprimera blanksteg
            import re
            usage_clean = re.sub(r'\s+', ' ', usage_raw).strip() if usage_raw else ''

            # Bygg en rik text för embedding
            text_parts = []
            text_parts.append(f"Frekvensområde: {freq_from} - {freq_to} MHz")
            if mangd:
                text_parts.append(f"Bandbredd: {mangd} MHz")
            if popular_name:
                text_parts.append(f"Populärnamn: {popular_name}")
            if duplex:
                text_parts.append(f"Duplex: {duplex}")
            if usage_clean:
                text_parts.append(f"Användning idag: {usage_clean}")
            if allocation_form:
                text_parts.append(f"Tilldelningsform: {allocation_form}")
            if planned_change:
                text_parts.append(f"Planerad förändring: {planned_change}")

            document_text = " | ".join(text_parts)

            # Metadata (viktigt för filtrering)
            meta = {
                "type": "allocation_plan",
                "source": str(file_path),
                "row": idx + 2,
                "frequency_from_mhz": freq_from_num,
                "frequency_to_mhz": freq_to_num,
                "frequency_from_str": freq_from,
                "frequency_to_str": freq_to,
                "bandwidth_mhz": mangd,
                "popular_name": popular_name,
                "duplex": duplex,
                "usage_today": usage_clean,
                "allocation_form": allocation_form,
                "planned_change": planned_change
            }
            items.append((document_text, meta))

        print(f"  📊 Extraherade {len(items)} objekt från spektrumplan {file_path.name}")
        return items

    # ------------------- Generisk fallback (en rad = ett dokument) -------------------
    def _extract_generic_csv_rows(self, df, file_path):
        items = []
        for idx, row in df.iterrows():
            # Gör om hela raden till en textsträng
            row_text = " | ".join([f"{col}: {str(val).strip()}" for col, val in row.items() if str(val).strip()])
            if row_text:
                meta = {
                    "type": "csv_generic",
                    "source": str(file_path),
                    "row": idx + 2
                }
                items.append((row_text, meta))
        return items

    # ------------------- ITU-typ (tidigare implementering, anpassad) -------------------
    def _extract_itu_csv_rows(self, df, file_path):
        import re
        items       = []
        rows_text   = []
        metas       = []

        # Identifiera kolumner (flexibelt)
        col_itu_band = next((c for c in df.columns if 'allokerat_band' in c), None)
        col_itu_services = next((c for c in df.columns if 'allokerade_tjänster' in c), None)
        col_swedish_usage = next((c for c in df.columns if 'användning' in c and not 'planerad' in c), None)
        col_freq_band = next((c for c in df.columns if 'frekvensband' in c), None)
        col_duplex = next((c for c in df.columns if 'duplexband' in c), None)
        col_anmärkning = next((c for c in df.columns if 'anmärkning' in c), None)

        if not col_swedish_usage:
            return items

        for idx, row in df.iterrows():
            # ITU-information (behåll som hel sträng, inklusive pipes)
            itu_band_full = row.get(col_itu_band, '').strip() if col_itu_band else ''
            itu_services_full = row.get(col_itu_services, '').strip() if col_itu_services else ''
            duplex = row.get(col_duplex, '').strip() if col_duplex else ''

            # Svenska användningar – dela på '|'
            usage_raw = row.get(col_swedish_usage, '').strip()
            if not usage_raw:
                continue
            swedish_usages = [u.strip() for u in usage_raw.split('|') if u.strip()]

            # Frekvensband – dela på '|' (kan vara fler eller färre än användningarna)
            freq_raw = row.get(col_freq_band, '').strip() if col_freq_band else ''
            freq_bands = [b.strip() for b in freq_raw.split('|') if b.strip()] if freq_raw else []

            # Anmärkningar – dela på '|'
            note_raw = row.get(col_anmärkning, '').strip() if col_anmärkning else ''
            notes = [n.strip() for n in note_raw.split('|') if n.strip()] if note_raw else []

            # Gör listorna lika långa som antalet svenska användningar
            num_usages = len(swedish_usages)
            # Fyll på freq_bands med tom sträng om för få
            if len(freq_bands) < num_usages:
                freq_bands += [''] * (num_usages - len(freq_bands))
            elif len(freq_bands) > num_usages:
                freq_bands = freq_bands[:num_usages]
            # Samma för notes
            if len(notes) < num_usages:
                notes += [''] * (num_usages - len(notes))
            elif len(notes) > num_usages:
                notes = notes[:num_usages]

            # Skapa ett dokument per svensk användning
            for usage_idx, (usage, band, note) in enumerate(zip(swedish_usages, freq_bands, notes)):
                text_parts = []
                if itu_band_full:
                    text_parts.append(f"ITU-band: {itu_band_full}")
                if itu_services_full:
                    text_parts.append(f"ITU-tjänster: {itu_services_full}")
                if band:
                    text_parts.append(f"Frekvensband: {band}")
                text_parts.append(f"Svensk användning: {usage}")
                if duplex:
                    text_parts.append(f"Duplex: {duplex}")
                if note:
                    # Rensa HTML-taggar om de finns (t.ex. <br/>)
                    note_clean = re.sub(r'<[^>]+>', ' ', note)
                    text_parts.append(f"Anmärkning: {note_clean}")

                document_text = " | ".join(text_parts)

                meta = {
                    "type": "itu_allocation",
                    "source": str(file_path),
                    "row": idx + 2,
                    "usage_index": usage_idx,
                    "itu_band": itu_band_full,
                    "itu_services": itu_services_full,
                    "swedish_usage": usage,
                    "frequency_band": band,
                    "duplex": duplex,
                    "anmärkning": note
                }
                items.append((document_text, meta))

        print(f"  📊 Extraherade {len(items)} rader från spektrumplan {file_path.name}")
        return items

    def _extract_excel_rows(self, file_path):
        items = []
        excel = pd.ExcelFile(file_path)
        for sheet_name in excel.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name, dtype=str).fillna('')
            for row_idx, row in df.iterrows():
                fields = []
                for col in df.columns:
                    val = str(row[col]).strip()
                    if val and val.lower() not in ['nan', 'none', '']:
                        fields.append(f"{col}: {val}")
                if fields:
                    text = " | ".join(fields)
                    meta = {"type": "excel_row", "source": str(file_path),
                            "sheet": sheet_name, "row": row_idx + 2}
                    items.append((text, meta))
        return items

    def _extract_pdf_tables_and_text(self, file_path):
        items = []
        # 1) Försök extrahera tabeller med tabula-py (kräver Java)
        try:
            import tabula
            tables = tabula.read_pdf(str(file_path), pages='all', multiple_tables=True)
            for table_idx, table in enumerate(tables):
                if table is None or table.empty:
                    continue
                for row_idx, row in table.iterrows():
                    fields = []
                    for col in table.columns:
                        val = str(row[col]).strip()
                        if val and val.lower() not in ['nan', 'none', '']:
                            fields.append(f"{col}: {val}")
                    if fields:
                        text = " | ".join(fields)
                        meta = {"type": "pdf_table_row", "source": str(file_path),
                                "table": table_idx, "row": row_idx + 1}
                        items.append((text, meta))
        except ImportError:
            print(f"   tabula-py ej installerad – hoppar över PDF-tabeller i {file_path.name}")
        except Exception as e:
            # Fångar även Java-fel
            print(f"   Kunde inte extrahera tabeller från PDF (Java saknas?): {e}")
            print(f"   Fortsätter med vanlig text från PDF.")

        # 2) Extrahera vanlig text med pypdf (alltid)
        text = self.read_pdf(file_path)
        if text and len(text) > 100:
            items.append((text, {"type": "pdf_text", "source": str(file_path)}))
        print(f"  📊 Extraherade {len(items)} rader från {file_path.name}")
        return items

    def _extract_docx_tables_and_text(self, file_path):
        items = []
        doc = Document(file_path)
        # Tabeller
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text = " | ".join(row_text)
                    meta = {"type": "docx_table_row", "source": str(file_path),
                            "table": table_idx, "row": row_idx + 1}
                    items.append((text, meta))
        # Löptext
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        if full_text:
            items.append((full_text, {"type": "docx_text", "source": str(file_path)}))
        return items

    def _extract_json_items(self, file_path):
        import json, re
        from html import unescape
        items = []

        with open(file_path, encoding='utf-8') as f:
            data = json.load(f)

        # Om data är en lista, iterera
        if isinstance(data, list):
            for obj in data:
                self._extract_spectrum_allocation_simple(obj, file_path, items)
        elif isinstance(data, dict):
            self._extract_spectrum_allocation_simple(data, file_path, items)
        else:
            # Generisk fallback (undantagsvis)
            self._extract_json_generic_fast(data, file_path, "", items)

        print(f"  📊 Extraherade {len(items)} inlägg från {file_path.name}")
        return items

    def _extract_spectrum_allocation_simple(self, obj, file_path, items):
        """Snabb extrahering av undantagsposter – skapar ett dokument per post."""
        import re
        from html import unescape

        # Hämta fält (hantera både med och utan .a-suffix)
        obj_id = obj.get('id', obj.get('id.a', ''))
        from_mhz = obj.get('fromMhz', obj.get('fromMhz.a', None))
        to_mhz = obj.get('toMhz', obj.get('toMhz.a', None))
        usage = obj.get('usageDescriptionLicenseExcemptionUsageDescription',
                        obj.get('usageDescriptionLicenseExcemptionUsageDescription.a', ''))
        reference = obj.get('reference', '')
        condition_raw = obj.get('condition', '')

        # Rensa condition snabbt (ta bort HTML-taggar, max 1500 tecken)
        condition_clean = ""
        if condition_raw and len(condition_raw) < 100000:  # undvik monsterfält
            # Ersätt <br> med newline
            temp = re.sub(r'<br\s*/?>', '\n', condition_raw, flags=re.IGNORECASE)
            # Ta bort alla andra taggar
            temp = re.sub(r'<[^>]+>', '', temp)
            # Avkoda entiteter (&auml; → ä)
            temp = unescape(temp)
            # Ta bort överflödiga blanksteg
            condition_clean = re.sub(r'\s+', ' ', temp).strip()
            # Trunkera om för långt
            if len(condition_clean) > 1500:
                condition_clean = condition_clean[:1500] + "..."

        # Bygg dokumenttext
        text_parts = []
        if obj_id:
            text_parts.append(f"ID: {obj_id}")
        if from_mhz is not None and to_mhz is not None:
            text_parts.append(f"Frekvensområde: {from_mhz} - {to_mhz} MHz")
        if usage:
            text_parts.append(f"Användning: {usage}")
        if reference:
            text_parts.append(f"Referens: {reference[:200]}")
        if condition_clean:
            text_parts.append(f"Villkor: {condition_clean}")

        if not text_parts:
            return  # inget att spara

        document_text = " | ".join(text_parts)
        meta = {
            "type": "spectrum_exemption",
            "source": str(file_path),
            "id": obj_id,
            "from_mhz": from_mhz,
            "to_mhz": to_mhz,
            "usage": usage,
            "reference": reference[:200]
        }
        items.append((document_text, meta))

    def _clean_html_fast(self, html_str):
        import re
        # Ersätt <br>, </p>, etc. med newline
        html_str = re.sub(r'<br\s*/?>', '\n', html_str, flags=re.IGNORECASE)
        # Ta bort alla andra taggar
        html_str = re.sub(r'<[^>]+>', '', html_str)
        # Ersätt HTML-entiteter (som &auml;)
        import html
        return html.unescape(html_str).strip()

    def _extract_json_generic_fast(self, data, file_path, prefix, items):
        """Rekursiv gång för övrig JSON, men med begränsning för att undvika explosion"""
        if len(items) > 10000:  # säkerhetsventil
            return
        if isinstance(data, dict):
            for k, v in data.items():
                new_prefix = f"{prefix}.{k}" if prefix else k
                self._extract_json_generic_fast(v, file_path, new_prefix, items)
        elif isinstance(data, list):
            for i, v in enumerate(data):
                new_prefix = f"{prefix}[{i}]" if prefix else f"[{i}]"
                self._extract_json_generic_fast(v, file_path, new_prefix, items)
        elif isinstance(data, str) and len(data) > 100:
            # Bara långa strängar, undvik korta
            items.append((data[:500], {"type": "json_text", "source": str(file_path), "key": prefix}))

    def _parse_freq_interval(self, interval_str):
        """Extraherar from/to från sträng som '3,8 < f ≤ 4,8 GHz' eller 'f ≤ 1,6 GHz'."""
        import re
        # Ta bort mellanslag och byt komma mot punkt
        s = interval_str.replace(',', '.').replace(' ', '')
        # Mönster för 'X < f ≤ Y' eller 'f ≤ Y' eller 'X < f'
        match = re.search(r'([\d\.]+)\s*[<≤]\s*f\s*[<≤]\s*([\d\.]+)', s)
        if match:
            return float(match.group(1)), float(match.group(2))
        match = re.search(r'f\s*[<≤]\s*([\d\.]+)', s)
        if match:
            return None, float(match.group(1))
        match = re.search(r'([\d\.]+)\s*[<≤]\s*f', s)
        if match:
            return float(match.group(1)), None
        return None, None

    def _extract_json_generic(self, data, file_path, prefix):
        """Fallback: rekursiv gång för andra JSON-strukturer."""
        items = []
        if isinstance(data, dict):
            for k, v in data.items():
                new_prefix = f"{prefix}.{k}" if prefix else k
                items.extend(self._extract_json_generic(v, file_path, new_prefix))
        elif isinstance(data, list):
            for i, v in enumerate(data):
                new_prefix = f"{prefix}[{i}]" if prefix else f"[{i}]"
                items.extend(self._extract_json_generic(v, file_path, new_prefix))
        elif isinstance(data, str) and len(data) > 50:
            # Bara lagom långa strängar, annars blir det för många dokument
            items.append((data, {"type": "json_text", "source": str(file_path), "key": prefix}))
        return items

    def _make_id(self, file_path, meta, idx):
        import hashlib
        unique = f"{file_path}_{meta.get('type')}_{meta.get('row', meta.get('chunk_index', idx))}"
        return hashlib.md5(unique.encode()).hexdigest()

    def save_metadata(self):
        with open(f"{self.CHROMA_DIR}/chromadb_params.json", "w") as f:
            json.dump({
                "COLLECTION_NAME"   : self.COLLECTION_NAME,
                "EMBED_MODEL"       : self.EMBED_MODEL,
                "DATA_ROOT"         : self.DATA_ROOT._str,
                "CHROMA_DIR"        : self.CHROMA_DIR,
                "BM25_DIR"          : self.BM25_DIR,
                "OLLAMA_HOST"       : self.OLLAMA_HOST,
                "EMBED_MODEL"       : self.EMBED_MODEL,
                "LLM_MODEL"         : self.LLM_MODEL,
                "BATCH_SIZE"        : self.BATCH_SIZE,
                "MAX_WORKERS"       : self.MAX_WORKERS
            }, f, indent=2)

        print(f"\n✅ KLART! {self.collection.count()} vektorer (en per dokument).")


if __name__=='__main__':

    rag = RAGSystem(data_root="./data", chroma_dir="./chroma_db")
    rag .create_chromadb_from_data( table_as_document=True )
    rag .save_metadata()
